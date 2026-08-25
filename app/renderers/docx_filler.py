import re
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def _soffice_cmd() -> str:
    for candidate in (
        "soffice",
        "libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if shutil.which(candidate) or Path(candidate).exists():
            return candidate
    raise RuntimeError("LibreOffice not found. Install LibreOffice to convert DOCX to PDF.")


def _iter_paragraphs(doc: Document):
    """Yield all paragraphs including those inside table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_in_paragraph(para, values: dict) -> None:
    """Replace {{placeholder}} tokens in a paragraph, handling cross-run splits."""
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    def sub(m: re.Match) -> str:
        return str(values.get(m.group(1).strip(), m.group(0)))

    replaced = PLACEHOLDER_RE.sub(sub, full_text)
    if replaced == full_text:
        return

    # Merge all runs into the first one, clear the rest.
    # This loses per-run formatting within a placeholder span — acceptable
    # because placeholders are always single, unformatted tokens.
    if para.runs:
        para.runs[0].text = replaced
        for run in para.runs[1:]:
            run.text = ""


def extract_placeholders_from_docx(docx_bytes: bytes) -> list[str]:
    doc = Document(BytesIO(docx_bytes))
    found: set[str] = set()
    for para in _iter_paragraphs(doc):
        full_text = "".join(run.text for run in para.runs)
        for m in PLACEHOLDER_RE.finditer(full_text):
            found.add(m.group(1))
    return sorted(found)


def fill_docx_placeholders(docx_bytes: bytes, values: dict) -> bytes:
    doc = Document(BytesIO(docx_bytes))
    for para in _iter_paragraphs(doc):
        _replace_in_paragraph(para, values)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    soffice = _soffice_cmd()
    with tempfile.TemporaryDirectory() as tmpdir:
        infile = Path(tmpdir) / "input.docx"
        infile.write_bytes(docx_bytes)

        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, str(infile)],
            capture_output=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr.decode(errors='replace')}"
            )

        outfile = Path(tmpdir) / "input.pdf"
        if not outfile.exists():
            raise RuntimeError("LibreOffice did not produce an output PDF")

        return outfile.read_bytes()
