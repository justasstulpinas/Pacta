import html as _html
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup, Tag
from io import BytesIO
import re


def _parse_color(color_str: str) -> RGBColor | None:
    if not color_str:
        return None
    # rgb(r, g, b)
    m = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", color_str)
    if m:
        return RGBColor(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    # #rrggbb
    m = re.match(r"#([0-9a-fA-F]{6})", color_str)
    if m:
        h = m.group(1)
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return None


def _parse_font_size(size_str: str) -> float | None:
    if not size_str:
        return None
    m = re.match(r"([\d.]+)(px|pt|em)?", size_str.strip())
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2) or "px"
    if unit == "px":
        return val * 0.75  # px → pt
    if unit == "pt":
        return val
    return None


def _parse_inline_styles(style_str: str) -> dict:
    result = {}
    if not style_str:
        return result
    for decl in style_str.split(";"):
        decl = decl.strip()
        if ":" not in decl:
            continue
        prop, _, val = decl.partition(":")
        result[prop.strip().lower()] = val.strip()
    return result


def _apply_run_styles(run, el: Tag, inherited: dict):
    """Apply inline CSS and semantic tags to a run."""
    styles = dict(inherited)

    # Merge inline style from current element
    if el.get("style"):
        styles.update(_parse_inline_styles(el["style"]))

    # Semantic tags walk up
    for parent in [el] + list(el.parents):
        if parent.name in ("strong", "b"):
            styles["font-weight"] = "bold"
        if parent.name in ("em", "i"):
            styles["font-style"] = "italic"
        if parent.name == "u":
            styles["text-decoration"] = "underline"
        if parent.name == "s":
            styles["text-decoration"] = "line-through"

    if styles.get("font-weight") in ("bold", "700"):
        run.bold = True
    if styles.get("font-style") == "italic":
        run.italic = True
    if "underline" in styles.get("text-decoration", ""):
        run.underline = True

    color = _parse_color(styles.get("color", ""))
    if color:
        run.font.color.rgb = color

    size = _parse_font_size(styles.get("font-size", ""))
    if size:
        run.font.size = Pt(size)

    font_family = styles.get("font-family", "")
    if font_family:
        # Take first font name
        name = font_family.split(",")[0].strip().strip("'\"")
        if name:
            run.font.name = name


def _add_runs_from_element(para, el, inherited_styles: dict):
    """Recursively walk element children and add runs."""
    if isinstance(el, str) or el.name is None:
        text = str(el) if isinstance(el, str) else el.string or ""
        if text:
            run = para.add_run(text)
            # apply inherited styles to plain text runs
            _apply_string_styles(run, inherited_styles)
        return

    # Collect styles for this element
    styles = dict(inherited_styles)
    if hasattr(el, "get") and el.get("style"):
        styles.update(_parse_inline_styles(el["style"]))
    if el.name in ("strong", "b"):
        styles["font-weight"] = "bold"
    if el.name in ("em", "i"):
        styles["font-style"] = "italic"
    if el.name == "u":
        styles["text-decoration"] = "underline"
    if el.name == "s":
        styles["text-decoration"] = "line-through"

    for child in el.children:
        if isinstance(child, str):
            text = child
            if text:
                run = para.add_run(text)
                _apply_string_styles(run, styles)
        elif hasattr(child, "name"):
            _add_runs_from_element(para, child, styles)


def _apply_string_styles(run, styles: dict):
    if styles.get("font-weight") in ("bold", "700"):
        run.bold = True
    if styles.get("font-style") == "italic":
        run.italic = True
    if "underline" in styles.get("text-decoration", ""):
        run.underline = True

    color = _parse_color(styles.get("color", ""))
    if color:
        run.font.color.rgb = color

    size = _parse_font_size(styles.get("font-size", ""))
    if size:
        run.font.size = Pt(size)

    font_family = styles.get("font-family", "")
    if font_family:
        name = font_family.split(",")[0].strip().strip("'\"")
        if name:
            run.font.name = name


def _set_para_alignment(para, style_str: str):
    if not style_str:
        return
    styles = _parse_inline_styles(style_str)
    align = styles.get("text-align", "")
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def render_docx_from_html(html: str) -> bytes:
    html = _html.unescape(html)
    soup = BeautifulSoup(html, "html.parser")
    document = Document()

    # Set default font
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    body = soup.find("body")
    if not body:
        return b""

    contract_body = body.find(class_="contract-body")
    content_el = contract_body if contract_body else body

    for el in content_el.children:
        if isinstance(el, str):
            text = el.strip()
            if text:
                para = document.add_paragraph()
                run = para.add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            continue

        if not hasattr(el, "name") or not el.name:
            continue

        tag = el.name

        if tag in ("h1", "h2", "h3"):
            level = int(tag[1])
            para = document.add_heading("", level=level)
            _set_para_alignment(para, el.get("style", ""))
            _add_runs_from_element(para, el, {})

        elif tag == "p":
            para = document.add_paragraph()
            _set_para_alignment(para, el.get("style", ""))
            _add_runs_from_element(para, el, {})

        elif tag in ("ul", "ol"):
            for li in el.find_all("li", recursive=False):
                para = document.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                _add_runs_from_element(para, li, {})

        elif tag == "blockquote":
            para = document.add_paragraph()
            para.style = "Quote"
            _add_runs_from_element(para, el, {})

        elif tag == "br":
            document.add_paragraph()

    # Footer
    footer_el = body.find(class_="footer")
    if footer_el:
        footer_text = footer_el.get_text(strip=True)
        for section in document.sections:
            fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
            fp.text = footer_text

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
