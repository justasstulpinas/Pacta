import io
import re
import shutil
import uuid as _uuid
from pathlib import Path
from fastapi import APIRouter, Depends, File, UploadFile, status
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

DOCX_UPLOAD_DIR = Path("app/uploads/templates")
DOCX_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

def _extract_placeholders(docx_bytes: bytes) -> list[str]:
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    found: set[str] = set()
    pattern = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
    def scan(text: str):
        for m in pattern.finditer(text):
            found.add(m.group(1))
    for para in doc.paragraphs:
        scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan(para.text)
    return sorted(found)

from app.database import get_db
from app.dependencies.auth import get_current_user
from app.models.user import User

from app.schemas.contract_template import (
    ContractTemplateCreate,
    ContractTemplateListItem,
    ContractTemplateOut,
    ContractTemplateUpdate
)
from app.schemas.filled_contract import FilledContractResponse

from app.services.contract_service import ContractService
from app.services.template_service import TemplateService


router = APIRouter(
    prefix="/templates",
    tags=["templates"],
)


@router.post(
    "",
    response_model=ContractTemplateOut,
    status_code=status.HTTP_201_CREATED,
)
def create_template(
    payload: ContractTemplateCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.create_template(payload, current_user)


@router.get(
    "",
    response_model=list[ContractTemplateListItem],
)
def list_templates(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    limit: int = 50,
    offset: int = 0,
    status: str | None = None,
    name: str | None = None,
):
    service = TemplateService(db)
    return service.list_user_templates(current_user, limit=limit, offset=offset, status=status, name=name)


@router.get(
    "/{template_id}",
    response_model=ContractTemplateOut,
)
def get_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.get_template_by_id(template_id, current_user)


@router.get(
    "/{template_id}/submissions",
    response_model=list[FilledContractResponse],
)
def get_submissions(
    template_id: int,
    limit: int = 20,
    offset: int = 0,
    status: str | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = ContractService(db)
    return service.get_template_submissions(
        template_id=template_id,
        user=current_user,
        limit=limit,
        offset=offset,
        status=status,
    )


@router.post(
    "/{template_id}/duplicate",
    response_model=ContractTemplateOut,
    status_code=status.HTTP_201_CREATED,
)
def duplicate_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.duplicate_template(template_id, current_user)


@router.patch(
    "/{template_id}/activate",
    response_model=ContractTemplateOut,
)
def activate_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.activate_template(template_id, current_user)


@router.patch(
    "/{template_id}/archive",
    response_model=ContractTemplateOut,
)
def archive_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.archive_template(template_id, current_user)


@router.delete(
    "/{template_id}",
    response_model=ContractTemplateOut,
)
def soft_delete_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.soft_delete_template(template_id, current_user)

@router.put(
    "/{template_id}",
    response_model=ContractTemplateOut,
)
def update_template(
    template_id: int,
    payload: ContractTemplateUpdate,
    db: Session =Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    service = TemplateService(db)
    return service.update_template(template_id, payload, current_user)


@router.post("/upload-docx")
async def upload_docx(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    if not (file.filename or "").lower().endswith(".docx"):
        from app.core.exceptions import BadRequestError
        raise BadRequestError("Tik .docx formato failai palaikomi")

    data = await file.read()
    file_key = f"tmp_{_uuid.uuid4().hex}"
    tmp_path = DOCX_UPLOAD_DIR / f"{file_key}.docx"
    tmp_path.write_bytes(data)

    placeholders = _extract_placeholders(data)
    original_name = (file.filename or "dokumentas").removesuffix(".docx")
    return {"file_key": file_key, "placeholders": placeholders, "filename": original_name}


class ReplaceTextRequest(BaseModel):
    file_key: str
    find_text: str
    placeholder: str


@router.post("/replace-text")
async def replace_text_in_docx(
    body: ReplaceTextRequest,
    current_user: User = Depends(get_current_user),
):
    from docx import Document
    from app.core.exceptions import BadRequestError, NotFoundError
    if not body.file_key.startswith("tmp_"):
        raise BadRequestError("Neleistinas failo raktas")
    path = DOCX_UPLOAD_DIR / f"{body.file_key}.docx"
    if not path.exists():
        raise NotFoundError("Failas nerastas")

    replacement = f"{{{{{body.placeholder}}}}}"
    doc = Document(str(path))

    def replace_in_para(para):
        full = "".join(r.text for r in para.runs)
        if body.find_text not in full:
            return
        new_full = full.replace(body.find_text, replacement, 1)
        # Put all text in first run, clear the rest
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(str(path))
    placeholders = _extract_placeholders(path.read_bytes())
    return {"placeholders": placeholders}


@router.get("/tmp/{file_key}")
async def get_tmp_docx(
    file_key: str,
    current_user: User = Depends(get_current_user),
):
    from app.core.exceptions import BadRequestError, NotFoundError
    if not file_key.startswith("tmp_"):
        raise BadRequestError("Neleistinas failo raktas")
    path = DOCX_UPLOAD_DIR / f"{file_key}.docx"
    if not path.exists():
        raise NotFoundError("Failas nerastas")
    return FileResponse(str(path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.get("/{template_id}/docx")
async def get_template_docx(
    template_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    from app.core.exceptions import NotFoundError, ForbiddenError
    service = TemplateService(db)
    template = service.get_template_by_id(template_id, current_user)
    if not template.docx_path:
        from app.core.exceptions import NotFoundError
        raise NotFoundError("DOCX failas nerastas")
    path = DOCX_UPLOAD_DIR / template.docx_path
    if not path.exists():
        raise NotFoundError("DOCX failas nerastas")
    return FileResponse(str(path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"{template.name}.docx")


def _docx_to_tiptap_html(data: bytes) -> str:
    import re
    import html as _html
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    _OPEN_TAG = re.compile(r'^<([a-zA-Z]+)(\s[^>]*)?>$')
    _CLOSE_TAG = re.compile(r'^</[a-zA-Z]+>$')
    _STYLE_ATTR = re.compile(r'style=["\']([^"\']*)["\']')
    _CSS_COLOR = re.compile(r'color\s*:\s*([^;]+)')
    _CSS_SIZE = re.compile(r'font-size\s*:\s*([^;]+)')
    _HTML_TAGS = re.compile(r'<[^>]+>')

    def _para_content(para) -> str:
        result = []
        pending: dict = {}  # styles from embedded opening tags

        for run in para.runs:
            text = run.text
            stripped = text.strip()

            # Detect embedded HTML opening tags (e.g. the run IS a <span style="...">)
            if _OPEN_TAG.match(stripped):
                style_m = _STYLE_ATTR.search(stripped)
                if style_m:
                    style_str = style_m.group(1).replace('&quot;', '"')
                    cm = _CSS_COLOR.search(style_str)
                    sm = _CSS_SIZE.search(style_str)
                    if cm:
                        pending['color'] = cm.group(1).strip()
                    if sm:
                        pending['font-size'] = sm.group(1).strip()
                continue

            # Detect embedded HTML closing tags
            if _CLOSE_TAG.match(stripped):
                pending.clear()
                continue

            # Real text — strip any remaining stray HTML tags
            clean = _HTML_TAGS.sub('', text)
            if not clean.strip():
                continue

            t = _html.escape(clean)

            # Apply Word run formatting
            if run.bold:
                t = f"<strong>{t}</strong>"
            if run.italic:
                t = f"<em>{t}</em>"
            if run.underline:
                t = f"<u>{t}</u>"

            # Apply accumulated inline styles from embedded tags
            if pending:
                style_str = '; '.join(f'{k}: {v}' for k, v in pending.items())
                t = f'<span style="{style_str}">{t}</span>'

            result.append(t)

        return ''.join(result)

    for para in doc.paragraphs:
        raw = _HTML_TAGS.sub('', para.text).strip()
        if not raw:
            parts.append("<p></p>")
            continue

        style_name = (para.style.name or "").lower()
        content = _para_content(para) or _html.escape(raw)

        align_style = ""
        try:
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                align_style = ' style="text-align: center;"'
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                align_style = ' style="text-align: right;"'
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                align_style = ' style="text-align: justify;"'
        except Exception:
            pass

        if "heading 1" in style_name or "title" in style_name:
            parts.append(f"<h1{align_style}>{content}</h1>")
        elif "heading 2" in style_name or "subtitle" in style_name:
            parts.append(f"<h2{align_style}>{content}</h2>")
        elif "heading 3" in style_name:
            parts.append(f"<h3{align_style}>{content}</h3>")
        elif "list" in style_name:
            parts.append(f"<li>{content}</li>")
        else:
            parts.append(f"<p{align_style}>{content}</p>")

    return "\n".join(parts)
